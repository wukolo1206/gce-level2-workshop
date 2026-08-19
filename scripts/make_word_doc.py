import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, title, text_list, fill_hex="E8F0FE", border_hex="1A73E8"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"💡 {title}")
    run_t.bold = True
    run_t.font.name = "Microsoft JhengHei"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0x15, 0x57, 0xB0)
    
    for item in text_list:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(f"• {item}")
        r.font.name = "Microsoft JhengHei"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x20, 0x21, 0x24)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_heading_1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(8)
    for r in h.runs:
        r.font.name = "Microsoft JhengHei"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    for r in h.runs:
        r.font.name = "Microsoft JhengHei"
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x20, 0x21, 0x24)
    return h

def add_body_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = "Microsoft JhengHei"
        r_pre.font.size = Pt(10.5)
        r_pre.font.color.rgb = RGBColor(0x20, 0x21, 0x24)
    r = p.add_run(text)
    r.font.name = "Microsoft JhengHei"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x3C, 0x40, 0x43)
    return p

def add_image_with_caption(doc, img_path, caption):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.8))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run(f"▲ {caption}")
        r_cap.font.name = "Microsoft JhengHei"
        r_cap.font.size = Pt(9.5)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

def main():
    doc = docx.Document()
    
    # 頁面邊界
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    img_dir = os.path.join(os.path.dirname(__file__), "images")
    
    # 主標題
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("Google Certified Educator Level 2 認證指南與講義")
    r_title.bold = True
    r_title.font.name = "Microsoft JhengHei"
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run("完全對照 Google 官方 Teacher Center 6 大核心單元與隨堂測驗原題 (最新圖文版)")
    r_sub.font.name = "Microsoft JhengHei"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)
    
    # 簡介呼應框
    add_callout_box(doc, "官方課程結構與備考對照說明", [
        "對照課程：Google for Education Level 2 (Intermediate use of Google Workspace)",
        "三大實戰模組：知識知能講義、精華速記卡、單頁 Web 互動刷題系統",
        "最新收錄重點：Gmail Auto-advance (自動進階)、Smart Canvas、Calendar Appointment Schedules"
    ])
    
    # 第一章
    add_heading_1(doc, "第一章：自動化課堂與行政任務 (Unit 1: Automate classroom tasks)")
    add_heading_2(doc, "1.1 Boost your efficiency in Gmail (提升 Gmail 郵件處理效率)")
    add_body_paragraph(doc, "在 Gmail 設定 (齒輪) -> 觀看所有設定 ->「進階」分頁中啟用【自動進階 (Auto-advance)】。當您刪除或封存郵件時，系統會自動直接呈現下一封未讀信件，大幅提升行政效率。", bold_prefix="★ 官方原題考點：Auto-advance (自動進階) — ")
    
    add_image_with_caption(doc, os.path.join(img_dir, "gmail_auto_advance.jpg"), "圖 1-1：Gmail 進階設定 - Auto-advance 自動進階功能介面示意圖")
    
    add_body_paragraph(doc, "在 Gmail 搜尋欄輸入 from:parent@school.edu 等條件，點選「建立篩選器」，可設定自動套用「家長來信」標籤或自動標示星號。", bold_prefix="郵件篩選器與標籤：")
    add_body_paragraph(doc, "預先撰寫罐頭回應範本，針對常見親師詢問一鍵帶入回覆。", bold_prefix="範本郵件 (Templates)：")
    
    add_heading_2(doc, "1.3 Level up collaboration with smart canvas (智慧畫布與智慧標籤)")
    add_body_paragraph(doc, "在 Google Docs/Sheets 中輸入 @ 符號，可帶出人員標籤 (@姓名)、檔案預閱卡片 (@檔名) 以及會議紀錄範本 (@Meeting notes)。", bold_prefix="Smart Chips (@ 智慧標籤)：")
    
    add_image_with_caption(doc, os.path.join(img_dir, "smart_canvas_chips.jpg"), "圖 1-2：Google Docs Smart Canvas 智慧標籤 (@ 符號功能) 介面示意圖")
    
    # 第二章
    add_heading_1(doc, "第二章：與家長及監護人高效溝通 (Unit 2: Communicate with parents)")
    add_heading_2(doc, "2.1 Organize guardian information with Google Forms")
    add_body_paragraph(doc, "收集家長 Email 或電話時，在簡答題右下角點選「驗證回應 (Data Validation)」，設定強制格式為「電子郵件位址」，防止誤填。", bold_prefix="資料驗證 (Data Validation)：")
    
    add_heading_2(doc, "2.3 Manage meetings with Google Workspace for Education")
    add_body_paragraph(doc, "在 Google Calendar 點選「建立」->「預約時間表 (Appointment Schedules)」，可自訂 15 分鐘諮詢時段，自動產生預約網址並連動 Meet 連結。", bold_prefix="預約時間表 (Appointment Schedules)：")
    
    add_image_with_caption(doc, os.path.join(img_dir, "calendar_appointments.jpg"), "圖 2-1：Google Calendar 預約時間表 (Appointment Schedules) 介面示意圖")
    
    # 第三章～第五章
    add_heading_1(doc, "第三章：系統化組織班級與教學素材 (Unit 3)")
    add_body_paragraph(doc, "選取文字點選「插入 -> 書籤」，於目錄文字設定超連結連至該書籤達成精準跳轉。", bold_prefix="Google Docs 書籤 (Bookmark)：")
    add_body_paragraph(doc, "學生製作數位學習歷程檔案，利用「在導覽列中隱藏 (Hide from navigation)」隱藏審查頁面。", bold_prefix="Google Sites 專題網頁：")
    
    add_heading_1(doc, "第四章：打造互動式自主學習環境 (Unit 4)")
    add_body_paragraph(doc, "利用形狀按鈕連結連至簡報內頁，將解答頁設定為「隱藏投影片 (Hide slide)」，防止點擊下一頁漏題。", bold_prefix="Google Slides 選擇板 (Choice Boards)：")
    
    add_heading_1(doc, "第五章：實施學生個人化與差異化學習 (Unit 5)")
    add_body_paragraph(doc, "發布作業時取消「所有學生」，指派給指定補救或資優小組。", bold_prefix="Classroom 差異化派發：")
    add_body_paragraph(doc, "單選題設定「根據回應跳轉至指定區段」，按答對答錯自動引導不同教材。", bold_prefix="Forms 適應性路徑：")
    
    # 第六章
    add_heading_1(doc, "第六章：分析與解讀學生學習數據 (Unit 6: Analyze & interpret data)")
    add_heading_2(doc, "6.2 Analyze data in Google Sheets")
    add_body_paragraph(doc, "=QUERY(A1:E100, \"SELECT A, B WHERE C < 60 AND E = '401'\", 1)", bold_prefix="QUERY SQL 語法：")
    add_body_paragraph(doc, "=IMPORTRANGE(\"URL\", \"工作表1!A1:D50\")。首次連線必須點選「允許存取 (Allow Access)」授權連線。", bold_prefix="IMPORTRANGE 跨表連線：")
    
    add_image_with_caption(doc, os.path.join(img_dir, "sheets_data_query.jpg"), "圖 6-1：Google Sheets QUERY 函數與樞紐分析數據視覺化介面示意圖")
    
    out_path = os.path.join(os.path.dirname(__file__), "Google_Certified_Educator_Level_2_講義.docx")
    doc.save(out_path)
    print("Successfully generated Word doc:", out_path)

if __name__ == "__main__":
    main()
